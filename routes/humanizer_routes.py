from io import BytesIO
from flask import Blueprint, render_template, request, send_file
from services.humanizer_service import humanize_text
import PyPDF2
import docx


humanizer_bp = Blueprint('humanizer', __name__)


def _read_uploaded_text(file_storage):
    if not file_storage or not file_storage.filename:
        return None

    filename = file_storage.filename.lower()
    content = file_storage.read()
    if not content:
        return None

    if filename.endswith('.pdf'):
        try:
            reader = PyPDF2.PdfReader(BytesIO(content))
            return '\n'.join(page.extract_text() or '' for page in reader.pages).strip()
        except Exception:
            return None

    if filename.endswith('.docx'):
        try:
            doc = docx.Document(BytesIO(content))
            return '\n'.join(paragraph.text for paragraph in doc.paragraphs).strip()
        except Exception:
            return None

    try:
        return content.decode('utf-8')
    except Exception:
        return content.decode('latin-1', errors='ignore')


@humanizer_bp.route('/humanizer', methods=['GET', 'POST'])
def humanizer():
    output = None
    orig_filename = None

    if request.method == 'POST':
        uploaded_file = request.files.get('file')
        user_text = _read_uploaded_text(uploaded_file) if uploaded_file else request.form.get('text')
        # preserve original filename when available (from upload or redirected form)
        orig_filename = (uploaded_file.filename if uploaded_file and uploaded_file.filename else request.form.get('orig_filename'))

        if user_text:
            output = humanize_text(user_text)
        else:
            output = 'Please paste text or upload a .txt, .pdf, or .docx file to humanize.'

    return render_template('tools/humanizer.html', output=output, orig_filename=orig_filename)


@humanizer_bp.route('/humanizer/download', methods=['POST'])
def humanizer_download():
    # Download the provided humanized content as a file. Supports .docx and .txt.
    content = request.form.get('download_content')
    filename = request.form.get('download_filename') or 'humanized.txt'

    if not content:
        return render_template('tools/humanizer.html', output='Nothing to download.' )

    filename = filename or 'humanized.txt'
    lower = filename.lower()
    if lower.endswith('.docx'):
        doc = docx.Document()
        for line in content.splitlines():
            doc.add_paragraph(line)
        bio = BytesIO()
        doc.save(bio)
        bio.seek(0)
        return send_file(bio, as_attachment=True, download_name=filename, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

    # default to plain text
    bio = BytesIO()
    bio.write((content or '').encode('utf-8'))
    bio.seek(0)
    return send_file(bio, as_attachment=True, download_name=filename, mimetype='text/plain')
